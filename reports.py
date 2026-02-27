"""
Report generators: Excel (.xlsx) and Word (.docx) export.
"""
import io
import datetime
import pandas as pd
import xlsxwriter


# ─── Color constants ───────────────────────────────────────────────────────────
CLR_HEADER     = "#1F3864"  # Dark blue header
CLR_SUBHEADER  = "#2E75B6"  # Medium blue
CLR_ALTO       = "#FCE4D6"  # Red light
CLR_MEDIO_ALTO = "#FFF0E0"  # Orange light
CLR_MEDIO      = "#FFFACD"  # Yellow light
CLR_BAJO       = "#F3F3F3"  # Gray
CLR_WHITE      = "#FFFFFF"
CLR_TEXT_HDR   = "#FFFFFF"
CLR_TOTAL_ROW  = "#D9E1F2"  # Light blue for totals


def _cop(v) -> str:
    try:
        return f"${float(v):,.0f}"
    except Exception:
        return str(v)


# ─── Excel Report ──────────────────────────────────────────────────────────────
def generate_excel(
    ventas: pd.DataFrame,
    compras: pd.DataFrame,
    kpis: dict,
    hallazgos: list[dict],
    iva_pivot: pd.DataFrame,
    empresa: str = "FAMIFAR,A VARIEDADES",
    nit: str = "1070951754",
    periodo: str = "Enero - Febrero 2026",
) -> bytes:
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})

    # ── Formats ───────────────────────────────────────────────────────────────
    fmt_title = workbook.add_format({
        "bold": True, "font_name": "Arial", "font_size": 14,
        "font_color": CLR_TEXT_HDR, "bg_color": CLR_HEADER,
        "align": "center", "valign": "vcenter", "border": 1,
    })
    fmt_header = workbook.add_format({
        "bold": True, "font_name": "Arial", "font_size": 10,
        "font_color": CLR_TEXT_HDR, "bg_color": CLR_HEADER,
        "align": "center", "valign": "vcenter", "border": 1, "text_wrap": True,
    })
    fmt_subheader = workbook.add_format({
        "bold": True, "font_name": "Arial", "font_size": 10,
        "font_color": CLR_TEXT_HDR, "bg_color": CLR_SUBHEADER,
        "align": "center", "valign": "vcenter", "border": 1,
    })
    fmt_data = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "valign": "vcenter",
    })
    fmt_money = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "num_format": '$#,##0', "align": "right",
    })
    fmt_total = workbook.add_format({
        "bold": True, "font_name": "Arial", "font_size": 10, "border": 2,
        "num_format": '$#,##0', "align": "right", "bg_color": CLR_TOTAL_ROW,
    })
    fmt_total_label = workbook.add_format({
        "bold": True, "font_name": "Arial", "font_size": 10,
        "bg_color": CLR_TOTAL_ROW, "border": 2,
    })
    fmt_risk_alto = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "bg_color": CLR_ALTO,
    })
    fmt_risk_medio_alto = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "bg_color": CLR_MEDIO_ALTO,
    })
    fmt_risk_medio = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "bg_color": CLR_MEDIO,
    })
    fmt_risk_bajo = workbook.add_format({
        "font_name": "Arial", "font_size": 9, "border": 1,
        "bg_color": CLR_BAJO,
    })

    def risk_fmt(nivel):
        if "ALTO" in nivel and "MEDIO" not in nivel:
            return fmt_risk_alto
        if "MEDIO-ALTO" in nivel:
            return fmt_risk_medio_alto
        if "MEDIO" in nivel:
            return fmt_risk_medio
        return fmt_risk_bajo

    now_str = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    # ── Sheet 1: Resumen Ejecutivo ─────────────────────────────────────────────
    ws1 = workbook.add_worksheet("📊 Resumen Ejecutivo")
    ws1.hide_gridlines(2)
    ws1.set_column("A:A", 30)
    ws1.set_column("B:C", 22)
    ws1.set_column("D:E", 18)
    ws1.set_row(0, 30)

    ws1.merge_range("A1:E1", f"REPORTE AUDITORÍA TRIBUTARIA — {empresa} | NIT {nit} | {periodo}", fmt_title)
    ws1.write("A2", f"Generado: {now_str}", workbook.add_format({"font_name": "Arial", "font_size": 8, "italic": True}))

    # KPI table
    kpi_hdr = workbook.add_format({**fmt_header.properties(), "bg_color": CLR_SUBHEADER})
    ws1.write_row(3, 0, ["KPI", "DIAN / Calculado", "Estado"], fmt_subheader)
    kpi_rows = [
        ("Total Ventas (DIAN)", kpis.get("total_ventas", 0), "Informativo"),
        ("Base Ventas (sin IVA)", kpis.get("base_ventas", 0), "Informativo"),
        ("IVA Generado", kpis.get("iva_generado", 0), "Verificar Form. 300"),
        ("Total Compras (DIAN)", kpis.get("total_compras", 0), "Informativo"),
        ("Base Compras (sin IVA)", kpis.get("base_compras", 0), "Informativo"),
        ("IVA Descontable", kpis.get("iva_descontable", 0), "Verificar Form. 300"),
        ("IVA Neto (Generado - Descontable)", kpis.get("iva_neto", 0), "A PAGAR" if kpis.get("iva_neto", 0) > 0 else "A FAVOR"),
        ("Margen Bruto Estimado %", kpis.get("margen_bruto", 0), "Ref. interna"),
        ("Facturas Ventas (FE)", kpis.get("num_facturas_ventas", 0), ""),
        ("Facturas Compras (FE)", kpis.get("num_facturas_compras", 0), ""),
        ("Notas Crédito Ventas", kpis.get("notas_credito_ventas", 0), "Revisar NC"),
        ("Notas Crédito Compras", kpis.get("notas_credito_compras", 0), "Revisar NC"),
    ]
    for i, (label, value, estado) in enumerate(kpi_rows, start=4):
        ws1.write(i, 0, label, fmt_data)
        if isinstance(value, float) and label.endswith("%"):
            ws1.write(i, 1, f"{value:.2f}%", fmt_data)
        elif isinstance(value, (int, float)):
            ws1.write(i, 1, value, fmt_money)
        else:
            ws1.write(i, 1, value, fmt_data)
        ws1.write(i, 2, estado, fmt_data)

    # Hallazgos summary table
    row = len(kpi_rows) + 6
    ws1.merge_range(row, 0, row, 4, "RESUMEN DE HALLAZGOS", fmt_title)
    row += 1
    ws1.write_row(row, 0, ["Código", "Nivel", "Área", "Descripción (resumen)", "Impacto COP"], fmt_header)
    row += 1
    for h in hallazgos:
        rf = risk_fmt(h["nivel"])
        ws1.write(row, 0, h["codigo"], rf)
        ws1.write(row, 1, h["nivel"], rf)
        ws1.write(row, 2, h["area"], rf)
        ws1.write(row, 3, h["descripcion"][:120], rf)
        ws1.write(row, 4, h["impacto"], fmt_money)
        row += 1

    # ── Sheet 2: Ventas DIAN ──────────────────────────────────────────────────
    ws2 = workbook.add_worksheet("📄 Ventas DIAN")
    ws2.hide_gridlines(2)
    _write_df_sheet(ws2, ventas, "VENTAS — FACTURAS ELECTRÓNICAS EMITIDAS",
                    empresa, nit, periodo, now_str,
                    workbook, fmt_title, fmt_header, fmt_data, fmt_money, fmt_total, fmt_total_label)

    # ── Sheet 3: Compras DIAN ─────────────────────────────────────────────────
    ws3 = workbook.add_worksheet("📦 Compras DIAN")
    ws3.hide_gridlines(2)
    _write_df_sheet(ws3, compras, "COMPRAS — FACTURAS ELECTRÓNICAS RECIBIDAS",
                    empresa, nit, periodo, now_str,
                    workbook, fmt_title, fmt_header, fmt_data, fmt_money, fmt_total, fmt_total_label)

    # ── Sheet 4: Hallazgos Detalle ────────────────────────────────────────────
    ws4 = workbook.add_worksheet("🔍 Hallazgos Detalle")
    ws4.hide_gridlines(2)
    ws4.set_column("A:A", 10)
    ws4.set_column("B:B", 18)
    ws4.set_column("C:C", 14)
    ws4.set_column("D:D", 60)
    ws4.set_column("E:E", 18)
    ws4.set_column("F:F", 14)
    ws4.set_column("G:G", 50)
    ws4.set_row(0, 28)
    ws4.merge_range("A1:G1", f"HALLAZGOS DE AUDITORÍA — {empresa} | {periodo}", fmt_title)
    ws4.write_row(1, 0, ["Código", "Nivel Riesgo", "Área", "Descripción", "Impacto COP", "Cuenta", "Norma Aplicable"], fmt_header)
    for i, h in enumerate(hallazgos, start=2):
        rf = risk_fmt(h["nivel"])
        ws4.write(i, 0, h["codigo"], rf)
        ws4.write(i, 1, h["nivel"], rf)
        ws4.write(i, 2, h["area"], rf)
        ws4.write(i, 3, h["descripcion"], rf)
        ws4.write(i, 4, h["impacto"], fmt_money)
        ws4.write(i, 5, h.get("cuenta", ""), rf)
        ws4.write(i, 6, h.get("norma", ""), rf)
        ws4.set_row(i, 40)

    # ── Sheet 5: Conciliación IVA ─────────────────────────────────────────────
    ws5 = workbook.add_worksheet("💰 Conciliación IVA")
    ws5.hide_gridlines(2)
    ws5.merge_range("A1:H1", f"CONCILIACIÓN IVA BIMESTRAL — {empresa} | {periodo}", fmt_title)
    ws5.write(2, 0, "⚠️ CASO DE ESTUDIO — VALORES ILUSTRATIVOS. Confrontar con Formulario 300 DIAN.",
              workbook.add_format({"font_name": "Arial", "font_size": 9, "italic": True, "font_color": "#C00000"}))
    if not iva_pivot.empty:
        cols = list(iva_pivot.columns)
        ws5.write_row(3, 0, cols, fmt_header)
        for r, row_data in enumerate(iva_pivot.itertuples(index=False), start=4):
            for c, val in enumerate(row_data):
                if isinstance(val, (int, float)):
                    ws5.write(r, c, val, fmt_money)
                else:
                    ws5.write(r, c, str(val), fmt_data)
        ws5.autofilter(3, 0, 3 + len(iva_pivot), len(cols) - 1)
    else:
        ws5.write(4, 0, "Sin datos suficientes para conciliación bimestral.", fmt_data)

    # ── Sheet 6: Procedimientos ───────────────────────────────────────────────
    ws6 = workbook.add_worksheet("📋 Procedimientos")
    ws6.hide_gridlines(2)
    ws6.set_column("A:A", 10)
    ws6.set_column("B:B", 14)
    ws6.set_column("C:C", 60)
    ws6.set_column("D:D", 20)
    ws6.set_column("E:E", 16)
    ws6.merge_range("A1:E1", f"PROGRAMA DE AUDITORÍA — {empresa} | {periodo}", fmt_title)
    ws6.write_row(1, 0, ["Hallazgo", "Área", "Procedimiento", "Responsable", "Estado"], fmt_header)
    for i, h in enumerate(hallazgos, start=2):
        rf = risk_fmt(h["nivel"])
        ws6.write(i, 0, h["codigo"], rf)
        ws6.write(i, 1, h["area"], rf)
        ws6.write(i, 2, h.get("procedimiento", ""), rf)
        ws6.write(i, 3, "Auditor", fmt_data)
        ws6.write(i, 4, "Pendiente", fmt_data)
        ws6.set_row(i, 50)

    workbook.close()
    return output.getvalue()


def _write_df_sheet(ws, df, titulo, empresa, nit, periodo, now_str,
                    workbook, fmt_title, fmt_header, fmt_data, fmt_money, fmt_total, fmt_total_label):
    """Helper to write a DataFrame to a worksheet with totals."""
    if df.empty:
        ws.write(0, 0, "Sin datos", fmt_data)
        return

    display_cols = [
        "Tipo de documento", "Folio", "Prefijo", "Fecha Emisión",
        "Nombre Emisor", "Nombre Receptor",
        "Base", "IVA", "Rete Renta", "Rete IVA", "Total", "Estado",
    ]
    available = [c for c in display_cols if c in df.columns]
    df_disp = df[available].copy()

    # Set column widths
    col_widths = {
        "Tipo de documento": 20, "Folio": 12, "Prefijo": 10,
        "Fecha Emisión": 14, "Nombre Emisor": 35, "Nombre Receptor": 35,
        "Base": 16, "IVA": 16, "Rete Renta": 14, "Rete IVA": 14,
        "Total": 18, "Estado": 12,
    }
    for ci, col in enumerate(available):
        ws.set_column(ci, ci, col_widths.get(col, 14))
    ws.set_row(0, 28)

    # Title
    ws.merge_range(0, 0, 0, len(available) - 1,
                   f"{titulo} — {empresa} | NIT {nit} | {periodo}", fmt_title)
    ws.write(1, 0, f"Generado: {now_str} | ⚠️ CASO DE ESTUDIO — VALORES ILUSTRATIVOS",
             workbook.add_format({"font_name": "Arial", "font_size": 8, "italic": True}))

    # Headers
    ws.write_row(2, 0, available, fmt_header)
    ws.freeze_panes(3, 0)
    ws.autofilter(2, 0, 2, len(available) - 1)

    # Data rows
    money_cols = {"Base", "IVA", "Rete Renta", "Rete IVA", "Total"}
    for ri, row in enumerate(df_disp.itertuples(index=False), start=3):
        for ci, (col, val) in enumerate(zip(available, row)):
            if col in money_cols and isinstance(val, (int, float)):
                ws.write(ri, ci, val, fmt_money)
            elif hasattr(val, 'strftime'):
                ws.write(ri, ci, str(val.date()) if pd.notna(val) else "", fmt_data)
            else:
                ws.write(ri, ci, "" if pd.isna(val) else val, fmt_data)

    # Totals row
    total_row = 3 + len(df_disp)
    ws.write(total_row, 0, "TOTAL", fmt_total_label)
    for ci, col in enumerate(available):
        if col in money_cols:
            ws.write(total_row, ci, df_disp[col].sum(), fmt_total)
        elif ci > 0:
            ws.write(total_row, ci, "", fmt_total_label)


# ─── Word Report ───────────────────────────────────────────────────────────────
def generate_word(
    kpis: dict,
    hallazgos: list[dict],
    empresa: str = "FAMIFAR,A VARIEDADES",
    nit: str = "1070951754",
    periodo: str = "Enero - Febrero 2026",
) -> bytes:
    """Generate a professional Word document with audit findings."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        return b""  # python-docx not available

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.replace("#", ""))
        tcPr.append(shd)

    def heading(doc, text, level=1, color_hex="1F3864"):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
        return p

    # ── Cover ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("REPORTE DE AUDITORÍA TRIBUTARIA")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{empresa}\nNIT: {nit}\nPeríodo: {periodo}\n"
                f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y')}").font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc = p.add_run("⚠️ CASO DE ESTUDIO — VALORES ILUSTRATIVOS")
    disc.bold = True
    disc.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_page_break()

    # ── Resumen ejecutivo ──────────────────────────────────────────────────────
    heading(doc, "1. RESUMEN EJECUTIVO", 1)
    kpi_data = [
        ("Total Ventas DIAN", _cop(kpis.get("total_ventas", 0))),
        ("Base Ventas (sin IVA)", _cop(kpis.get("base_ventas", 0))),
        ("IVA Generado", _cop(kpis.get("iva_generado", 0))),
        ("Total Compras DIAN", _cop(kpis.get("total_compras", 0))),
        ("Base Compras (sin IVA)", _cop(kpis.get("base_compras", 0))),
        ("IVA Descontable", _cop(kpis.get("iva_descontable", 0))),
        ("IVA Neto", _cop(kpis.get("iva_neto", 0))),
        ("Margen Bruto Estimado", f"{kpis.get('margen_bruto', 0):.1f}%"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Indicador"
    hdr[1].text = "Valor"
    for cell in hdr:
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for label, value in kpi_data:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    # ── Hallazgos ──────────────────────────────────────────────────────────────
    heading(doc, "2. HALLAZGOS DE AUDITORÍA", 1)
    risk_colors = {
        "🔴 ALTO": "FCE4D6",
        "🟠 MEDIO-ALTO": "FFF0E0",
        "🟡 MEDIO": "FFFACD",
        "🟣 MEDIO": "F3E5F5",
        "⚪ BAJO-MEDIO": "F3F3F3",
    }
    for h in hallazgos:
        heading(doc, f"{h['codigo']} — {h['nivel']} | {h['area']}", 2, "2E75B6")
        doc.add_paragraph(h["descripcion"])
        t2 = doc.add_table(rows=4, cols=2)
        t2.style = "Table Grid"
        fields = [
            ("Cuenta Contable", h.get("cuenta", "")),
            ("Impacto Estimado COP", _cop(h.get("impacto", 0))),
            ("Norma Aplicable", h.get("norma", "")),
            ("Procedimiento de Auditoría", h.get("procedimiento", "")),
        ]
        bg = risk_colors.get(h["nivel"], "FFFFFF")
        for ri, (k, v) in enumerate(fields):
            row = t2.rows[ri].cells
            row[0].text = k
            row[1].text = v
            set_cell_bg(row[0], "1F3864")
            set_cell_bg(row[1], bg)
            for para in row[0].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
        doc.add_paragraph()

    # ── Conclusiones ──────────────────────────────────────────────────────────
    heading(doc, "3. CONCLUSIONES Y RECOMENDACIONES", 1)
    altos = [h for h in hallazgos if "ALTO" in h["nivel"] and "MEDIO" not in h["nivel"]]
    doc.add_paragraph(
        f"Del análisis de {len(hallazgos)} hallazgo(s) identificados en el período {periodo}, "
        f"{len(altos)} corresponde(n) a riesgo ALTO que requieren atención inmediata. "
        "Se recomienda conciliar la información con los formularios 300 (IVA) y 350 (Retenciones) "
        "presentados ante la DIAN, y ajustar los saldos contables si hay inconsistencias."
    )
    doc.add_paragraph(
        "⚠️ Este reporte es un caso de estudio con valores ilustrativos. "
        "Los hallazgos deben ser verificados contra la contabilidad oficial y declaraciones tributarias."
    )

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
